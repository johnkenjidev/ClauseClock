"""ClauseClock Stage 1 backend tests: auth, isolation, ingestion, provenance,
scanned detection, and hard delete via GridFS."""
import io
import os
import uuid

import pytest
import requests
from dotenv import load_dotenv
load_dotenv("/app/frontend/.env")
load_dotenv("/app/backend/.env")
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

BASE_URL = os.environ["REACT_APP_BACKEND_URL"].rstrip("/")
API = f"{BASE_URL}/api"

SEED_EMAIL = "test@clauseclock.app"
SEED_PASSWORD = "Test1234!"

SCANNED_MSG = (
    "This looks like a scanned or image-based PDF. ClauseClock cannot read it "
    "yet. Upload a text-based version."
)


# ---------- helpers ----------
def make_pdf(text: str = "This is a real contract. " * 40, pages: int = 1) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    for p in range(pages):
        y = 750
        for line in (f"{text} page-{p+1}").split(". "):
            c.drawString(50, y, line[:90])
            y -= 14
            if y < 50:
                break
        c.showPage()
    c.save()
    return buf.getvalue()


def make_blank_pdf(pages: int = 1) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    for _ in range(pages):
        c.showPage()  # no text
    c.save()
    return buf.getvalue()


def make_docx() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph("This amendment covers renewal and pricing terms.")
    doc.add_heading("Section 1 - Renewal", level=2)
    doc.add_paragraph("The Agreement auto-renews for 12 months unless notice is given.")
    doc.add_paragraph("Notice period is 60 days before renewal date.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def new_session():
    return requests.Session()


def register_or_login(session, email, password="Test1234!"):
    r = session.post(f"{API}/auth/register", json={"email": email, "password": password})
    if r.status_code == 400:
        r = session.post(f"{API}/auth/login", json={"email": email, "password": password})
    assert r.status_code == 200, f"auth failed: {r.status_code} {r.text}"
    return r.json()


# ---------- fixtures ----------
@pytest.fixture(scope="module")
def seed_session():
    s = new_session()
    r = s.post(f"{API}/auth/login", json={"email": SEED_EMAIL, "password": SEED_PASSWORD})
    assert r.status_code == 200, f"seed login failed: {r.text}"
    return s


@pytest.fixture(scope="module")
def user_a():
    email = f"test_a_{uuid.uuid4().hex[:8]}@clauseclock.app"
    s = new_session()
    register_or_login(s, email)
    return s, email


@pytest.fixture(scope="module")
def user_b():
    email = f"test_b_{uuid.uuid4().hex[:8]}@clauseclock.app"
    s = new_session()
    register_or_login(s, email)
    return s, email


# ---------- Auth ----------
class TestAuth:
    def test_register_login_me(self):
        email = f"test_reg_{uuid.uuid4().hex[:8]}@clauseclock.app"
        s = new_session()
        r = s.post(f"{API}/auth/register", json={"email": email, "password": "Test1234!"})
        assert r.status_code == 200
        data = r.json()
        assert data["email"] == email
        assert "id" in data
        # cookie should be set
        assert "access_token" in s.cookies.get_dict()

        # /me works
        r = s.get(f"{API}/auth/me")
        assert r.status_code == 200
        assert r.json()["email"] == email

        # logout clears session
        r = s.post(f"{API}/auth/logout")
        assert r.status_code == 200

        # login again
        s2 = new_session()
        r = s2.post(f"{API}/auth/login", json={"email": email, "password": "Test1234!"})
        assert r.status_code == 200
        r = s2.get(f"{API}/auth/me")
        assert r.status_code == 200

    def test_no_session_returns_401(self):
        s = new_session()
        r = s.get(f"{API}/contracts")
        assert r.status_code == 401

    def test_invalid_credentials(self):
        s = new_session()
        r = s.post(f"{API}/auth/login", json={"email": SEED_EMAIL, "password": "wrong"})
        assert r.status_code == 401


# ---------- Contracts / ingestion ----------
class TestPDFIngestion:
    def test_upload_pdf_creates_contract(self, seed_session):
        pdf = make_pdf(pages=2)
        files = {"file": ("agreement.pdf", pdf, "application/pdf")}
        data = {"name": "TEST_ACME MSA", "counterparty": "Acme Inc",
                "doc_role": "primary", "annual_value": "24000", "currency": "USD"}
        r = seed_session.post(f"{API}/contracts", files=files, data=data)
        assert r.status_code == 200, r.text
        body = r.json()
        c = body["contract"]
        d = body["document"]
        assert c["name"] == "TEST_ACME MSA"
        assert c["counterparty"] == "Acme Inc"
        assert c["annual_value"] == 24000
        assert c["currency"] == "USD"
        assert c["value_source"] == "user_entered"
        # provenance nulls for user_entered
        assert c["value_source_quote"] is None
        assert c["value_source_document_id"] is None
        assert c["value_source_location"] is None
        assert c["value_source_char_offset"] is None
        # doc
        assert d["extraction_method"] == "pdfplumber"
        assert d["sha256"] and len(d["sha256"]) == 64
        assert d["size_bytes"] > 0
        assert d["page_count"] and d["page_count"] > 0
        assert d["raw_text"]
        assert "Page 1" in d["raw_text"] or "Page" in d["raw_text"]
        assert d["doc_role"] == "primary"
        # cleanup
        seed_session.delete(f"{API}/contracts/{c['id']}")

    def test_contract_without_annual_value(self, seed_session):
        pdf = make_pdf()
        files = {"file": ("no_value.pdf", pdf, "application/pdf")}
        data = {"name": "TEST_no_value", "doc_role": "primary"}
        r = seed_session.post(f"{API}/contracts", files=files, data=data)
        assert r.status_code == 200
        c = r.json()["contract"]
        assert c["annual_value"] is None
        assert c["value_source"] is None
        seed_session.delete(f"{API}/contracts/{c['id']}")


class TestDocxIngestion:
    def test_add_docx_amendment(self, seed_session):
        # create primary first
        files = {"file": ("primary.pdf", make_pdf(), "application/pdf")}
        r = seed_session.post(f"{API}/contracts", files=files,
                              data={"name": "TEST_docx_parent", "doc_role": "primary"})
        assert r.status_code == 200
        cid = r.json()["contract"]["id"]

        # add docx amendment
        files = {"file": ("amend.docx", make_docx(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = seed_session.post(f"{API}/contracts/{cid}/documents", files=files,
                              data={"doc_role": "amendment"})
        assert r.status_code == 200, r.text
        d = r.json()["document"]
        assert d["extraction_method"] == "python-docx"
        assert d["doc_role"] == "amendment"
        rt = d["raw_text"]
        assert rt
        assert "[§" in rt and "¶" in rt

        seed_session.delete(f"{API}/contracts/{cid}")


class TestScannedDetection:
    def test_blank_pdf_detected_as_scanned(self, seed_session):
        pdf = make_blank_pdf(pages=1)
        files = {"file": ("scanned.pdf", pdf, "application/pdf")}
        data = {"name": "TEST_scanned", "doc_role": "primary"}
        r = seed_session.post(f"{API}/contracts", files=files, data=data)
        assert r.status_code == 200
        body = r.json()
        d = body["document"]
        assert d["extraction_method"] == "failed_no_text"
        assert SCANNED_MSG in d["extraction_warnings"]
        seed_session.delete(f"{API}/contracts/{body['contract']['id']}")


class TestIsolation:
    def test_user_a_cannot_access_user_b_contract(self, user_a, user_b):
        sa, _ = user_a
        sb, _ = user_b

        # user A creates contract
        files = {"file": ("a.pdf", make_pdf(), "application/pdf")}
        r = sa.post(f"{API}/contracts", files=files,
                    data={"name": "TEST_A_secret", "doc_role": "primary"})
        assert r.status_code == 200
        cid = r.json()["contract"]["id"]

        # user B cannot GET
        r = sb.get(f"{API}/contracts/{cid}")
        assert r.status_code == 404

        # user B cannot DELETE
        r = sb.delete(f"{API}/contracts/{cid}")
        assert r.status_code == 404

        # user B list does not include A's contract
        r = sb.get(f"{API}/contracts")
        assert r.status_code == 200
        ids = [c["id"] for c in r.json()["contracts"]]
        assert cid not in ids

        # user A cleanup
        sa.delete(f"{API}/contracts/{cid}")

    def test_user_id_from_client_ignored(self, seed_session):
        # attempt to pass user_id as form param — should be ignored (attaches to session)
        files = {"file": ("x.pdf", make_pdf(), "application/pdf")}
        data = {"name": "TEST_userid_spoof", "doc_role": "primary",
                "user_id": "000000000000000000000000"}
        r = seed_session.post(f"{API}/contracts", files=files, data=data)
        assert r.status_code == 200
        c = r.json()["contract"]
        # verify it's owned by session user by getting it back
        r = seed_session.get(f"{API}/contracts/{c['id']}")
        assert r.status_code == 200
        seed_session.delete(f"{API}/contracts/{c['id']}")


class TestContractDetail:
    def test_get_contract_returns_documents(self, seed_session):
        files = {"file": ("cd.pdf", make_pdf(), "application/pdf")}
        r = seed_session.post(f"{API}/contracts", files=files,
                              data={"name": "TEST_detail", "doc_role": "primary"})
        cid = r.json()["contract"]["id"]

        r = seed_session.get(f"{API}/contracts/{cid}")
        assert r.status_code == 200
        body = r.json()
        assert body["contract"]["id"] == cid
        assert len(body["documents"]) == 1
        assert body["documents"][0]["raw_text"]
        assert body["documents"][0]["doc_role"] == "primary"

        seed_session.delete(f"{API}/contracts/{cid}")


class TestHardDelete:
    def test_delete_removes_gridfs_originals(self, seed_session):
        # count fs.files before via mongo directly (using motor via env)
        import pymongo
        mc = pymongo.MongoClient(os.environ["MONGO_URL"])
        db = mc[os.environ["DB_NAME"]]

        before = db["fs.files"].count_documents({})

        files = {"file": ("hd.pdf", make_pdf(), "application/pdf")}
        r = seed_session.post(f"{API}/contracts", files=files,
                              data={"name": "TEST_harddelete", "doc_role": "primary"})
        cid = r.json()["contract"]["id"]

        # add another doc so we can verify multi-file delete
        files = {"file": ("hd2.docx", make_docx(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = seed_session.post(f"{API}/contracts/{cid}/documents", files=files,
                              data={"doc_role": "amendment"})
        assert r.status_code == 200

        after_upload = db["fs.files"].count_documents({})
        assert after_upload == before + 2

        r = seed_session.delete(f"{API}/contracts/{cid}")
        assert r.status_code == 200
        assert r.json()["deleted"] is True

        # verify 404
        r = seed_session.get(f"{API}/contracts/{cid}")
        assert r.status_code == 404

        # verify docs collection empty for this contract
        assert db.documents.count_documents({"contract_id": cid}) == 0

        # verify GridFS files removed
        after_delete = db["fs.files"].count_documents({})
        assert after_delete == before, f"GridFS not cleaned: before={before} after={after_delete}"

        mc.close()
