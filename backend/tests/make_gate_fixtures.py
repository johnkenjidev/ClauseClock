"""Generate Stage-1 gate fixtures: 5 difficult text-based contracts + 1 scanned."""
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                Spacer, Table, TableStyle, PageBreak, SimpleDocTemplate)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from docx import Document
from PIL import Image, ImageDraw

OUT = os.path.dirname(os.path.abspath(__file__))
styles = getSampleStyleSheet()
body = styles["BodyText"]
h1 = styles["Heading1"]; h2 = styles["Heading2"]

CLAUSE = ("This Agreement shall automatically renew for successive twelve (12) "
          "month terms unless either party provides written notice not less than "
          "sixty (60) days before the end of the then-current term. All notices "
          "shall be sent by certified mail to the General Counsel at 123 Main St.")

# 1) DOCX with headings, paragraphs, tables, appendix ----------------------
def make_docx():
    d = Document()
    d.add_heading("Master Services Agreement", level=1)
    d.add_paragraph("This Agreement is entered into between Acme Corp and Vendor Inc.")
    d.add_heading("8. Term and Termination", level=2)
    d.add_paragraph(CLAUSE)
    d.add_heading("9. Fees", level=2)
    t = d.add_table(rows=3, cols=2); t.style = "Table Grid"
    t.cell(0,0).text="Item"; t.cell(0,1).text="Amount"
    t.cell(1,0).text="Annual License"; t.cell(1,1).text="$24,000"
    t.cell(2,0).text="Support"; t.cell(2,1).text="$3,600"
    d.add_heading("Appendix A — Service Levels", level=2)
    d.add_paragraph("Uptime commitment of 99.9% measured monthly. Service credits "
                    "of 5% apply per 0.1% shortfall, claimable within thirty (30) days.")
    d.save(os.path.join(OUT, "gate_docx.docx"))

# 2) Two-column PDF --------------------------------------------------------
def make_twocol():
    path = os.path.join(OUT, "gate_twocol.pdf")
    doc = BaseDocTemplate(path, pagesize=letter,
                          leftMargin=0.6*inch, rightMargin=0.6*inch,
                          topMargin=0.7*inch, bottomMargin=0.7*inch)
    gap = 0.3*inch
    fw = (doc.width - gap) / 2
    f1 = Frame(doc.leftMargin, doc.bottomMargin, fw, doc.height, id="c1")
    f2 = Frame(doc.leftMargin + fw + gap, doc.bottomMargin, fw, doc.height, id="c2")
    doc.addPageTemplates([PageTemplate(id="two", frames=[f1, f2])])
    story = [Paragraph("MASTER SERVICES AGREEMENT (Two Column)", h1)]
    for i in range(1, 40):
        story.append(Paragraph(f"Section {i}. " + CLAUSE, body))
        story.append(Spacer(1, 6))
    doc.build(story)

# 3) 40+ page PDF ----------------------------------------------------------
def make_long():
    path = os.path.join(OUT, "gate_40pages.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    story = [Paragraph("LONG SERVICES AGREEMENT", h1)]
    for p in range(1, 46):
        story.append(Paragraph(f"Article {p}", h2))
        for _ in range(6):
            story.append(Paragraph(CLAUSE, body)); story.append(Spacer(1, 6))
        story.append(PageBreak())
    doc.build(story)

# 4) Appendix PDF (key terms only in the appendix at the end) --------------
def make_appendix():
    path = os.path.join(OUT, "gate_appendix.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    story = [Paragraph("SUPPLY AGREEMENT", h1)]
    for p in range(1, 8):
        story.append(Paragraph(f"Clause {p}. General provisions and boilerplate text.", body))
        story.append(Spacer(1, 6))
    story.append(PageBreak())
    story.append(Paragraph("APPENDIX B — Pricing and Renewal", h2))
    story.append(Paragraph("The annual contract value is $48,000. " + CLAUSE, body))
    story.append(Paragraph("Price increases are capped at eight percent (8%) per renewal.", body))
    doc.build(story)

# 5) Tables PDF ------------------------------------------------------------
def make_tables():
    path = os.path.join(OUT, "gate_tables.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    data = [["Tier", "Annual Fee", "Notice (days)", "Cap %"],
            ["Bronze", "$12,000", "30", "5%"],
            ["Silver", "$24,000", "60", "8%"],
            ["Gold", "$48,000", "90", "10%"]]
    tbl = Table(data, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey)]))
    story = [Paragraph("PRICING SCHEDULE", h1), Spacer(1, 12), tbl,
             Spacer(1, 18), Paragraph(CLAUSE, body)]
    doc.build(story)

# 6) Scanned/image-only PDF (no text layer) --------------------------------
def make_scanned():
    img = Image.new("RGB", (1000, 1300), "white")
    dr = ImageDraw.Draw(img)
    dr.text((60, 60), "MASTER SERVICES AGREEMENT (scanned copy)", fill="black")
    dr.text((60, 120), CLAUSE[:80], fill="black")
    img_path = os.path.join(OUT, "_scan.png")
    img.save(img_path)
    path = os.path.join(OUT, "gate_scanned.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    c.drawImage(img_path, 40, 120, width=520, height=650)
    c.showPage(); c.save()
    os.remove(img_path)

if __name__ == "__main__":
    make_docx(); make_twocol(); make_long(); make_appendix(); make_tables(); make_scanned()
    for f in sorted(os.listdir(OUT)):
        if f.startswith("gate_"):
            print(f, os.path.getsize(os.path.join(OUT, f)), "bytes")
